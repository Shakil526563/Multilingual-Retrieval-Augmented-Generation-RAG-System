import os
import re
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path

import fitz  # PyMuPDF
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import faiss
import pickle

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS

from groq import Groq
from django.conf import settings

logger = logging.getLogger(__name__)


class BengaliTextProcessor:
    """Handles Bengali and English text preprocessing and cleaning"""
    
    @staticmethod
    def clean_bengali_text(text: str) -> str:
        """Clean and normalize Bengali text with improved Unicode handling"""
        # Remove extra whitespaces
        text = re.sub(r'\s+', ' ', text)
        
        # Remove unwanted characters but preserve Bengali script and common symbols
        # Bengali Unicode range: \u0980-\u09FF
        # Keep Bengali digits, punctuation, and common English characters
        text = re.sub(r'[^\u0980-\u09FF\u0020-\u007E\s\u2013\u2014\u2018\u2019\u201C\u201D]', '', text)
        
        # Fix common Bengali text issues
        text = re.sub(r'[\u09BC]+', '\u09BC', text)  # Normalize nukta
        text = re.sub(r'[\u09CD\u09CD]+', '\u09CD', text)  # Normalize hasanta
        
        # Remove multiple punctuation
        text = re.sub(r'[।]{2,}', '।', text)
        text = re.sub(r'[?]{2,}', '?', text)
        text = re.sub(r'[!]{2,}', '!', text)
        
        # Fix spacing around Bengali punctuation
        text = re.sub(r'\s+([।!?])', r'\1', text)
        text = re.sub(r'([।!?])\s*', r'\1 ', text)
        
        return text.strip()
    
    @staticmethod
    def clean_english_text(text: str) -> str:
        """Clean and normalize English text"""
        # Remove extra whitespaces
        text = re.sub(r'\s+', ' ', text)
        
        # Remove unwanted characters but preserve English letters, numbers, and common punctuation
        text = re.sub(r'[^\w\s.,!?;:\-\'"()\[\]{}]', '', text)
        
        # Remove multiple punctuation marks
        text = re.sub(r'[.]{3,}', '...', text)  
        text = re.sub(r'[!]{2,}', '!', text)    
        text = re.sub(r'[?]{2,}', '?', text)   
        text = re.sub(r'[-]{2,}', '--', text)   
        
        # Clean up spacing around punctuation
        text = re.sub(r'\s+([.,!?;:])', r'\1', text)  
        text = re.sub(r'([.,!?;:])\s*', r'\1 ', text)  
        
        # Fix common formatting issues
        text = re.sub(r'(\w)\s*\'\s*(\w)', r"\1'\2", text)  # Fix contractions like "don ' t" -> "don't"
        text = re.sub(r'\s*\(\s*', ' (', text)  # Fix spacing around parentheses
        text = re.sub(r'\s*\)\s*', ') ', text)
        
        return text.strip()
    
    @staticmethod
    def clean_mixed_text(text: str) -> str:
        """Clean text that may contain both Bengali and English with font-aware processing"""
        # Remove extra whitespaces
        text = re.sub(r'\s+', ' ', text)
        
  
        text = re.sub(r'[^\u0980-\u09FF\w\s.,!?;:\-\'"()\[\]{}।\u2013\u2014\u2018\u2019\u201C\u201D]', '', text)
        
        # Fix Bengali-specific issues from PDF extraction
        text = re.sub(r'[\u09BC]+', '\u09BC', text) 
        text = re.sub(r'[\u09CD]+', '\u09CD', text)  
        
        # Remove multiple punctuation for both languages
        text = re.sub(r'[।]{2,}', '।', text)  
        text = re.sub(r'[.]{3,}', '...', text) 
        text = re.sub(r'[!]{2,}', '!', text)    
        text = re.sub(r'[?]{2,}', '?', text)    
        
        # Clean up spacing around punctuation (both Bengali and English)
        text = re.sub(r'\s+([.,!?;:।])', r'\1', text)  
        text = re.sub(r'([.,!?;:।])\s*', r'\1 ', text)  
        
        # Fix common PDF extraction artifacts for Bengali text
        text = re.sub(r'\s+([কখগঘঙচছজঝঞটঠডঢণতথদধনপফবভমযরলশষসহ])', r' \1', text)
        
        return text.strip()
    
    @staticmethod
    def preprocess_text(text: str) -> str:
        """Automatically detect language and apply appropriate cleaning"""
        language = BengaliTextProcessor.detect_language(text)
        
        if language == 'bengali':
            return BengaliTextProcessor.clean_bengali_text(text)
        elif language == 'english':
            return BengaliTextProcessor.clean_english_text(text)
        else:
            # For mixed or unknown text, use mixed cleaning
            return BengaliTextProcessor.clean_mixed_text(text)
    
    @staticmethod
    def detect_language(text: str) -> str:
        """Detect if text is Bengali or English"""
        bengali_chars = len(re.findall(r'[\u0980-\u09FF]', text))
        english_chars = len(re.findall(r'[a-zA-Z]', text))
        
        if bengali_chars > english_chars:
            return 'bengali'
        elif english_chars > 0:
            return 'english'
        else:
            return 'unknown'


class DocumentProcessor:
    """Handles document processing and text extraction for PDF, TXT, and DOCX files"""
    
    def __init__(self):
        self.text_processor = BengaliTextProcessor()
        self.bengali_font_path = self._get_bengali_font_path()
    
    def _get_bengali_font_path(self) -> Optional[str]:
        """Get the path to the Bengali font file"""
        try:
            # Get the base directory (project root)
            base_dir = Path(__file__).parent.parent
            font_path = base_dir / "fronts" / "Siyamrupali.ttf"
            
            if font_path.exists():
                logger.info(f"Found Bengali font at: {font_path}")
                return str(font_path)
            else:
                logger.warning(f"Bengali font not found at: {font_path}")
                return None
        except Exception as e:
            logger.error(f"Error locating Bengali font: {str(e)}")
            return None
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from various file formats (PDF, TXT, DOCX)"""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.txt':
            return self.extract_text_from_txt(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def extract_text_from_txt(self, txt_path: str) -> str:
        """Extract text from TXT file with Bengali support"""
        try:
            logger.info(f"Opening TXT file: {txt_path}")
            
            # Try different encodings for better Bengali support
            encodings_to_try = ['utf-8', 'utf-8-sig', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings_to_try:
                try:
                    with open(txt_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    
                    # Check if Bengali characters are present and properly decoded
                    bengali_chars = len(re.findall(r'[\u0980-\u09FF]', text))
                    logger.info(f"Using encoding {encoding}: found {bengali_chars} Bengali characters")
                    
                    # Clean the extracted text
                    cleaned_text = self.text_processor.preprocess_text(text)
                    
                    logger.info(f"Total extracted text length: {len(cleaned_text)} characters")
                    return cleaned_text
                    
                except UnicodeDecodeError:
                    logger.warning(f"Failed to decode with {encoding}, trying next encoding")
                    continue
            
            # If all encodings fail, raise an error
            raise ValueError(f"Could not decode text file {txt_path} with any supported encoding")
            
        except Exception as e:
            logger.error(f"Error extracting text from TXT {txt_path}: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file with Bengali support"""
        try:
            logger.info(f"Opening DOCX file: {docx_path}")
            
            # Import python-docx if available
            try:
                from docx import Document
            except ImportError:
                logger.error("python-docx package not installed. Install with: pip install python-docx")
                return ""
            
            doc = Document(docx_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            # Clean the extracted text
            cleaned_text = self.text_processor.preprocess_text(text)
            
            logger.info(f"Total extracted text length: {len(cleaned_text)} characters")
            return cleaned_text
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {docx_path}: {str(e)}")
            return ""

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF file using PyMuPDF with Bengali font support"""
        try:
            logger.info(f"Opening PDF file: {pdf_path}")
            doc = fitz.open(pdf_path)
            text = ""
            
            logger.info(f"PDF has {len(doc)} pages")
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Try multiple text extraction methods for better Bengali support
                page_text = self._extract_text_with_font_support(page)
                
                logger.info(f"Page {page_num + 1} extracted {len(page_text)} characters")
                
                # Clean the extracted text
                cleaned_text = self.text_processor.preprocess_text(page_text)
                text += cleaned_text + "\n"
            
            doc.close()
            logger.info(f"Total extracted text length: {len(text)} characters")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF {pdf_path}: {str(e)}")
            return ""
    
    def _extract_text_with_font_support(self, page) -> str:
        """Extract text with Bengali font support using multiple methods"""
        text = ""
        
        try:
            # Method 1: Standard text extraction
            standard_text = page.get_text()
            
            # Method 2: Text extraction with font information
            text_dict = page.get_text("dict")
            font_aware_text = self._extract_from_text_dict(text_dict)
            
            # Method 3: Text extraction with layout preservation
            layout_text = page.get_text("layout")
            
            # Choose the best extraction based on Bengali character count
            texts = [standard_text, font_aware_text, layout_text]
            bengali_counts = []
            
            for t in texts:
                bengali_chars = len(re.findall(r'[\u0980-\u09FF]', t))
                bengali_counts.append(bengali_chars)
            
            # Use the text with the highest Bengali character count
            best_index = bengali_counts.index(max(bengali_counts))
            text = texts[best_index]
            
            logger.debug(f"Selected extraction method {best_index + 1} with {bengali_counts[best_index]} Bengali characters")
            
        except Exception as e:
            logger.warning(f"Error in font-aware extraction, falling back to standard: {str(e)}")
            text = page.get_text()
        
        return text
    
    def _extract_from_text_dict(self, text_dict: Dict) -> str:
        """Extract text from PyMuPDF text dictionary with font awareness"""
        text = ""
        
        try:
            for block in text_dict.get("blocks", []):
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        line_text = ""
                        for span in line.get("spans", []):
                            span_text = span.get("text", "")
                            font_name = span.get("font", "").lower()
                            
                            # Check if this span might contain Bengali text
                            has_bengali = bool(re.search(r'[\u0980-\u09FF]', span_text))
                            
                            # Log font information for Bengali text
                            if has_bengali:
                                logger.debug(f"Bengali text found with font: {font_name}")
                            
                            line_text += span_text
                        
                        text += line_text + "\n"
                    text += "\n"
        except Exception as e:
            logger.warning(f"Error extracting from text dict: {str(e)}")
        
        return text
    
    def chunk_document(self, text: str, chunk_size: int = 1000, chunk_overlap: int = 100) -> List[Document]:
        """Split document into chunks for better retrieval"""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", "।", ".", "!", "?", " ", ""]
        )
        
        chunks = text_splitter.split_text(text)
        documents = []
        
        for i, chunk in enumerate(chunks):
            if chunk.strip():  # Only add non-empty chunks
                doc = Document(
                    page_content=chunk.strip(),
                    metadata={
                        'chunk_id': i,
                        'language': self.text_processor.detect_language(chunk),
                        'chunk_size': len(chunk)
                    }
                )
                documents.append(doc)
        
        return documents
    
    def analyze_pdf_content(self, pdf_path: str) -> Dict[str, Any]:
        """Analyze PDF content to understand text distribution and fonts"""
        try:
            doc = fitz.open(pdf_path)
            analysis = {
                'total_pages': len(doc),
                'bengali_pages': 0,
                'english_pages': 0,
                'mixed_pages': 0,
                'fonts_used': set(),
                'total_bengali_chars': 0,
                'total_english_chars': 0
            }
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                
                # Analyze text content
                bengali_chars = len(re.findall(r'[\u0980-\u09FF]', page_text))
                english_chars = len(re.findall(r'[a-zA-Z]', page_text))
                
                analysis['total_bengali_chars'] += bengali_chars
                analysis['total_english_chars'] += english_chars
                
                # Categorize page
                if bengali_chars > english_chars and bengali_chars > 0:
                    analysis['bengali_pages'] += 1
                elif english_chars > bengali_chars and english_chars > 0:
                    analysis['english_pages'] += 1
                elif bengali_chars > 0 or english_chars > 0:
                    analysis['mixed_pages'] += 1
                
                # Extract font information
                try:
                    text_dict = page.get_text("dict")
                    for block in text_dict.get("blocks", []):
                        if block.get("type") == 0:  # Text block
                            for line in block.get("lines", []):
                                for span in line.get("spans", []):
                                    font_name = span.get("font", "")
                                    if font_name:
                                        analysis['fonts_used'].add(font_name)
                except Exception as e:
                    logger.debug(f"Could not extract font info from page {page_num}: {e}")
            
            doc.close()
            
            # Convert set to list for JSON serialization
            analysis['fonts_used'] = list(analysis['fonts_used'])
            
            logger.info(f"PDF Analysis: {analysis['bengali_pages']} Bengali pages, "
                       f"{analysis['english_pages']} English pages, "
                       f"{analysis['mixed_pages']} mixed pages")
            
            return analysis
            
        except Exception as e:
            logger.error(f"Error analyzing PDF content: {str(e)}")
            return {}


class VectorStore:
    """Manages vector embeddings and similarity search"""
    
    def __init__(self, model_name: str = "sentence-transformers/all-MiniLM-L6-v2"):
        self.embeddings = HuggingFaceEmbeddings(model_name=model_name)
        self.vector_store = None
        self.documents = []
        
    def create_vector_store(self, documents: List[Document]) -> None:
        """Create FAISS vector store from documents"""
        if not documents:
            raise ValueError("No documents provided to create vector store")
        
        self.documents = documents
        self.vector_store = FAISS.from_documents(documents, self.embeddings)
        logger.info(f"Created vector store with {len(documents)} documents")
    
    def save_vector_store(self, path: str) -> None:
        """Save vector store to disk"""
        if self.vector_store is None:
            raise ValueError("No vector store to save")
        
        os.makedirs(path, exist_ok=True)
        self.vector_store.save_local(path)
        
        # Save documents separately
        with open(os.path.join(path, 'documents.pkl'), 'wb') as f:
            pickle.dump(self.documents, f)
        
        logger.info(f"Vector store saved to {path}")
    
    def load_vector_store(self, path: str) -> None:
        """Load vector store from disk"""
        if not os.path.exists(path):
            raise FileNotFoundError(f"Vector store path {path} does not exist")
        
        try:
            # Try to load with allow_dangerous_deserialization parameter (newer versions)
            self.vector_store = FAISS.load_local(path, self.embeddings, allow_dangerous_deserialization=True)
        except TypeError:
            # Fallback for older FAISS versions that don't support the parameter
            self.vector_store = FAISS.load_local(path, self.embeddings)
        
        # Load documents separately
        docs_path = os.path.join(path, 'documents.pkl')
        if os.path.exists(docs_path):
            with open(docs_path, 'rb') as f:
                self.documents = pickle.load(f)
        
        logger.info(f"Vector store loaded from {path}")
    
    def similarity_search(self, query: str, k: int = 5) -> List[Document]:
        """Perform similarity search"""
        if self.vector_store is None:
            raise ValueError("Vector store not initialized")
        
        return self.vector_store.similarity_search(query, k=k)
    
    def similarity_search_with_score(self, query: str, k: int = 5) -> List[tuple]:
        """Perform similarity search with scores"""
        if self.vector_store is None:
            raise ValueError("Vector store not initialized")
        
        return self.vector_store.similarity_search_with_score(query, k=k)


class ConversationMemory:
    """Manages short-term and long-term memory"""
    
    def __init__(self, max_short_term: int = 10):
        self.short_term_memory = []  # Recent chat history
        self.max_short_term = max_short_term
        self.long_term_memory = None  # Vector store with documents
    
    def add_to_short_term(self, user_query: str, assistant_response: str) -> None:
        """Add conversation to short-term memory"""
        self.short_term_memory.append({
            'user': user_query,
            'assistant': assistant_response,
            'timestamp': np.datetime64('now')
        })
        
        # Keep only recent conversations
        if len(self.short_term_memory) > self.max_short_term:
            self.short_term_memory = self.short_term_memory[-self.max_short_term:]
    
    def get_context_from_memory(self) -> str:
        """Get formatted context from short-term memory"""
        if not self.short_term_memory:
            return ""
        
        context = "Previous conversation context:\n"
        for entry in self.short_term_memory[-3:]:  # Last 3 exchanges
            context += f"User: {entry['user']}\n"
            context += f"Assistant: {entry['assistant']}\n\n"
        
        return context
    
    def set_long_term_memory(self, vector_store: VectorStore) -> None:
        """Set the vector store as long-term memory"""
        self.long_term_memory = vector_store


class GroqLLMClient:
    """Handles interaction with Groq API"""
    
    def __init__(self, api_key: str):
        if not api_key:
            raise ValueError("Groq API key not provided")
        self.client = Groq(api_key=api_key)
        self.model = "llama3-8b-8192"  # Fast model for better response times
    
    def generate_response(self, prompt: str, max_tokens: int = 1000) -> str:
        """Generate response using Groq API"""
        try:
            chat_completion = self.client.chat.completions.create(
                messages=[
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                model=self.model,
                max_tokens=max_tokens,
                temperature=0.7,
            )
            
            return chat_completion.choices[0].message.content
        except Exception as e:
            logger.error(f"Error generating response: {str(e)}")
            return "I apologize, but I'm having trouble generating a response right now."


class RAGEvaluator:
    """Evaluates RAG system performance"""
    
    def __init__(self):
        self.metrics = {
            'groundedness_scores': [],
            'relevance_scores': [],
            'response_times': []
        }
    
    def evaluate_groundedness(self, response: str, retrieved_docs: List[Document]) -> float:
        """Evaluate if response is grounded in retrieved context"""
        if not retrieved_docs:
            return 0.0
        
        # Simple keyword-based groundedness check
        response_words = set(response.lower().split())
        context_words = set()
        
        for doc in retrieved_docs:
            context_words.update(doc.page_content.lower().split())
        
        if not context_words:
            return 0.0
        
        overlap = len(response_words.intersection(context_words))
        groundedness = overlap / len(response_words) if response_words else 0.0
        
        self.metrics['groundedness_scores'].append(groundedness)
        return groundedness
    
    def evaluate_relevance(self, query: str, retrieved_docs: List[Document]) -> float:
        """Evaluate relevance of retrieved documents to query"""
        if not retrieved_docs:
            return 0.0
        
        # Use embeddings to calculate semantic similarity
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        
        query_embedding = embeddings.embed_query(query)
        doc_embeddings = embeddings.embed_documents([doc.page_content for doc in retrieved_docs])
        
        similarities = cosine_similarity([query_embedding], doc_embeddings)[0]
        avg_relevance = np.mean(similarities)
        
        self.metrics['relevance_scores'].append(avg_relevance)
        return avg_relevance
    
    def get_evaluation_summary(self) -> Dict[str, float]:
        """Get summary of evaluation metrics"""
        summary = {}
        
        if self.metrics['groundedness_scores']:
            summary['avg_groundedness'] = np.mean(self.metrics['groundedness_scores'])
        
        if self.metrics['relevance_scores']:
            summary['avg_relevance'] = np.mean(self.metrics['relevance_scores'])
        
        if self.metrics['response_times']:
            summary['avg_response_time'] = np.mean(self.metrics['response_times'])
        
        return summary


class MultilingualRAGSystem:
    """Main RAG system class that orchestrates all components"""
    
    def __init__(self):
        self.document_processor = DocumentProcessor()
        self.vector_store = VectorStore()
        self.memory = ConversationMemory()
        self.llm_client = None  # Will be initialized when needed
        self.evaluator = RAGEvaluator()
        self.text_processor = BengaliTextProcessor()
        self.is_initialized = False
    
    def _ensure_llm_client(self):
        """Ensure LLM client is initialized"""
        if self.llm_client is None:
            api_key = getattr(settings, 'GROQ_API_KEY', None)
            if not api_key:
                raise ValueError("GROQ_API_KEY not configured in settings")
            self.llm_client = GroqLLMClient(api_key)
    
    def initialize_system(self, pdf_path: str, force_rebuild: bool = False) -> None:
        """Initialize the RAG system with PDF documents"""
        vector_db_path = str(settings.VECTOR_DB_PATH)
        
        # Try to load existing vector store
        if not force_rebuild and os.path.exists(vector_db_path):
            try:
                self.vector_store.load_vector_store(vector_db_path)
                self.memory.set_long_term_memory(self.vector_store)
                self.is_initialized = True
                logger.info("Loaded existing vector store")
                return
            except Exception as e:
                logger.warning(f"Failed to load existing vector store: {e}")
        
        # Process document and create new vector store
        logger.info(f"Processing document: {pdf_path}")
        
        # First, analyze the document content (if it's a PDF)
        if pdf_path.lower().endswith('.pdf'):
            pdf_analysis = self.document_processor.analyze_pdf_content(pdf_path)
            if pdf_analysis:
                logger.info(f"PDF contains {pdf_analysis.get('total_bengali_chars', 0)} Bengali characters "
                           f"and {pdf_analysis.get('total_english_chars', 0)} English characters")
                if pdf_analysis.get('fonts_used'):
                    logger.info(f"Fonts found in PDF: {', '.join(pdf_analysis['fonts_used'][:5])}")
        
        # Extract text with enhanced support for multiple formats
        text = self.document_processor.extract_text_from_file(pdf_path)
        
        if not text.strip():
            raise ValueError("No text extracted from document")
        
        # Chunk the document
        documents = self.document_processor.chunk_document(text)
        logger.info(f"Created {len(documents)} document chunks")
        
        # Log language distribution in chunks
        language_count = {}
        for doc in documents:
            lang = doc.metadata.get('language', 'unknown')
            language_count[lang] = language_count.get(lang, 0) + 1
        logger.info(f"Document chunk languages: {language_count}")
        
        # Create and save vector store
        self.vector_store.create_vector_store(documents)
        self.vector_store.save_vector_store(vector_db_path)
        
        # Set up memory
        self.memory.set_long_term_memory(self.vector_store)
        self.is_initialized = True
        logger.info("RAG system initialized successfully")
    
    def query(self, user_query: str, k: int = 5) -> Dict[str, Any]:
        """Process user query and return response with metadata"""
        if not self.is_initialized:
            raise ValueError("RAG system not initialized")
        
        # Ensure LLM client is ready
        self._ensure_llm_client()
        
        start_time = np.datetime64('now')
        
        # Detect query language
        query_language = self.text_processor.detect_language(user_query)
        
        # Retrieve relevant documents
        retrieved_docs = self.vector_store.similarity_search_with_score(user_query, k=k)
        documents = [doc for doc, score in retrieved_docs]
        scores = [score for doc, score in retrieved_docs]
        
        # Get conversation context
        conversation_context = self.memory.get_context_from_memory()
        
        # Prepare context for LLM
        context = ""
        for i, doc in enumerate(documents):
            context += f"Document {i+1} (Similarity: {1-scores[i]:.3f}):\n{doc.page_content}\n\n"
        
        # Create prompt
        prompt = self._create_prompt(user_query, context, conversation_context, query_language)
        
        # Generate response
        response = self.llm_client.generate_response(prompt)
        
        # Add to memory
        self.memory.add_to_short_term(user_query, response)
        
        # Evaluate response
        groundedness = self.evaluator.evaluate_groundedness(response, documents)
        relevance = self.evaluator.evaluate_relevance(user_query, documents)
        
        # Calculate response time
        end_time = np.datetime64('now')
        response_time = float((end_time - start_time) / np.timedelta64(1, 's'))
        self.evaluator.metrics['response_times'].append(response_time)
        
        return {
            'response': response,
            'query_language': query_language,
            'retrieved_documents': [
                {
                    'content': doc.page_content[:200] + '...' if len(doc.page_content) > 200 else doc.page_content,
                    'similarity_score': 1 - scores[i],
                    'metadata': doc.metadata
                }
                for i, doc in enumerate(documents)
            ],
            'evaluation': {
                'groundedness': groundedness,
                'relevance': relevance,
                'response_time': response_time
            }
        }
    
    def _create_prompt(self, query: str, context: str, conversation_context: str, language: str) -> str:
        """Create a well-structured prompt for the LLM"""
        language_instruction = {
            'bengali': "Please respond in Bengali (বাংলা) language.",
            'english': "Please respond in English language.",
            'unknown': "Please respond in the same language as the query."
        }.get(language, "Please respond appropriately.")
        
        prompt = f"""
You are a highly intelligent and helpful AI assistant capable of understanding and answering both Bengali and English queries. Use the provided context documents to generate accurate, grounded, and concise responses.

🔹 INSTRUCTIONS:
1. Always prioritize information found in the context documents.
2. If the user asks "who" or "কাকে", and the answer is a person's name, clearly return only the **name**/"নাম" in your response.
3. If the user's question is in Bengali, answer in Bengali. If the question is in English, answer in English.
4. If the context does NOT contain enough information to answer the question reliably, clearly say so.
5. Be concise but complete — include names, numbers, and details when directly supported by the context.
6. Maintain consistency and continuity with previous queries if relevant.

{conversation_context}

CONTEXT DOCUMENTS:
{context}

USER QUESTION: {query}

ANSWER:"""
        
        return prompt
    
    def get_system_stats(self) -> Dict[str, Any]:
        """Get system statistics and evaluation metrics"""
        stats = {
            'is_initialized': self.is_initialized,
            'total_documents': len(self.vector_store.documents) if self.vector_store.documents else 0,
            'conversation_history_length': len(self.memory.short_term_memory),
            'evaluation_summary': self.evaluator.get_evaluation_summary()
        }
        
        if self.vector_store.documents:
            languages = [doc.metadata.get('language', 'unknown') for doc in self.vector_store.documents]
            language_counts = {lang: languages.count(lang) for lang in set(languages)}
            stats['document_languages'] = language_counts
        
        return stats
    
    def analyze_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """Analyze PDF file and return detailed information about content and fonts"""
        return self.document_processor.analyze_pdf_content(pdf_path)
    
    def analyze_document(self, file_path: str) -> Dict[str, Any]:
        """Analyze document file and return detailed information about content"""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.document_processor.analyze_pdf_content(file_path)
        else:
            # For non-PDF files, provide basic analysis
            try:
                text = self.document_processor.extract_text_from_file(file_path)
                bengali_chars = len(re.findall(r'[\u0980-\u09FF]', text))
                english_chars = len(re.findall(r'[a-zA-Z]', text))
                
                return {
                    'file_type': file_extension,
                    'total_characters': len(text),
                    'total_bengali_chars': bengali_chars,
                    'total_english_chars': english_chars,
                    'language': 'bengali' if bengali_chars > english_chars else 'english' if english_chars > 0 else 'unknown'
                }
            except Exception as e:
                logger.error(f"Error analyzing document: {str(e)}")
                return {}
    
    def get_bengali_font_info(self) -> Dict[str, Any]:
        """Get information about the Bengali font configuration"""
        font_info = {
            'font_path': self.document_processor.bengali_font_path,
            'font_available': self.document_processor.bengali_font_path is not None
        }
        
        if font_info['font_available']:
            try:
                font_path = Path(self.document_processor.bengali_font_path)
                font_info.update({
                    'font_name': font_path.name,
                    'font_size': font_path.stat().st_size,
                    'font_exists': font_path.exists()
                })
            except Exception as e:
                font_info['error'] = str(e)
        
        return font_info
